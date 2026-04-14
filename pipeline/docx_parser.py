"""
pipeline/docx_parser.py
-----------------------
DOCX → List[{page_num, title, text, char_count}]

Dùng python-docx để extract text theo từng heading section.
Mỗi heading = 1 "page" (tương đương slide trong PDF).

Install: pip install python-docx
"""

from pathlib import Path


# ── Constants ──────────────────────────────────────────────────────
MIN_CHARS = 50

# Heading styles của Word
HEADING_STYLES = {
    "Heading 1", "Heading 2", "Heading 3",
    "Title", "Subtitle",
    # Tiếng Việt
    "Tieu de 1", "Tieu de 2",
}


def parse_docx(file_path: str) -> list[dict]:
    """
    Parse DOCX → List[page_dict] tương thích với pdf_parser output.

    Strategy: gom các paragraph liên tiếp vào 1 "section"
    mỗi khi gặp Heading → tạo section mới.
    """
    try:
        from docx import Document
    except ImportError:
        raise ImportError("Chua cai python-docx. Chay: pip install python-docx")

    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"File khong ton tai: {file_path}")

    doc = Document(file_path)
    sections = []
    current_title = path.stem   # fallback title = tên file
    current_lines = []
    page_num = 0

    def flush_section():
        nonlocal page_num
        text = "\n".join(current_lines).strip()
        if len(text) >= MIN_CHARS:
            page_num += 1
            sections.append({
                "page_num":   page_num,
                "title":      current_title,
                "text":       text,
                "char_count": len(text),
            })

    for para in doc.paragraphs:
        style_name = para.style.name if para.style else ""
        text       = para.text.strip()

        if not text:
            continue

        is_heading = (
            style_name in HEADING_STYLES
            or style_name.startswith("Heading")
            or style_name.startswith("Title")
        )

        if is_heading:
            # Flush section hiện tại trước khi bắt đầu section mới
            if current_lines:
                flush_section()
                current_lines = []
            current_title = text
        else:
            current_lines.append(text)

    # Flush section cuối
    if current_lines:
        flush_section()

    # Nếu không có heading nào → toàn bộ doc là 1 section
    if not sections:
        full_text = "\n".join(
            p.text.strip() for p in doc.paragraphs if p.text.strip()
        )
        if full_text:
            sections.append({
                "page_num":   1,
                "title":      path.stem,
                "text":       full_text,
                "char_count": len(full_text),
            })

    return sections


def get_docx_metadata(file_path: str) -> dict:
    """Metadata cơ bản của DOCX."""
    try:
        from docx import Document
        doc = Document(file_path)
        core = doc.core_properties
        return {
            "total_pages": len([p for p in doc.paragraphs if p.text.strip()]),
            "title":       core.title or "",
            "author":      core.author or "",
            "file_size_kb": round(Path(file_path).stat().st_size / 1024, 1),
        }
    except Exception:
        return {"total_pages": 0, "title": "", "author": "", "file_size_kb": 0}


# ── CLI test ───────────────────────────────────────────────────────
if __name__ == "__main__":
    import sys
    if sys.platform == "win32":
        sys.stdout.reconfigure(encoding="utf-8", errors="replace")

    if len(sys.argv) < 2:
        print("Usage: python -m pipeline.docx_parser <file.docx>")
        sys.exit(1)

    pages = parse_docx(sys.argv[1])
    print(f"\n[OK] Parsed {len(pages)} sections\n")
    print(f"{'NUM':<5} {'TITLE':<45} {'CHARS'}")
    print("-" * 60)
    for p in pages[:10]:
        print(f"  {p['page_num']:<3} {p['title'][:43]:<45} {p['char_count']}")
    if len(pages) > 10:
        print(f"  ... ({len(pages)-10} more)")
    if pages:
        print(f"\n--- Sample text (section 1) ---")
        print(pages[0]["text"][:300])